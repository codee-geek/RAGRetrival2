from __future__ import annotations

import hashlib
import os
import tempfile
from datetime import datetime, timezone
from pathlib import Path

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.services.bm25 import get_bm25_indexer, get_session_bm25_path
from app.services.embeddings import get_embeddings
from app.services.local_storage import sanitize_session_id, save_document
from app.services.pinecone_store import upsert_chunks
from core.config import CHUNK_OVERLAP, CHUNK_SIZE

SUPPORTED_EXTENSIONS = [".pdf", ".txt", ".docx"]

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    separators=["\n\n", "\n", ".", " "],
)


def calculate_content_hash(file_bytes: bytes) -> str:
    """Calculate MD5 hash of uploaded bytes for stable document ids."""
    return hashlib.md5(file_bytes).hexdigest()


def _load_docx(file_path: str):
    """Load a .docx as plain text via python-docx (avoids the heavy unstructured dep)."""
    import docx

    document = docx.Document(file_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)
    return [Document(page_content=text, metadata={"source": file_path})]


def load_document(file_path: str):
    """Load document based on file extension."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return PyPDFLoader(file_path).load()
    elif ext == ".txt":
        return TextLoader(file_path, encoding="utf-8").load()
    elif ext == ".docx":
        return _load_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _write_temp_file(filename: str, file_bytes: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
        tmp_file.write(file_bytes)
        return tmp_file.name


def _prepare_chunks(
    *,
    documents,
    filename: str,
    session_id: str,
    document_id: str,
    upload_time: str,
):
    # semantic_section labels each chunk with its source document name.
    semantic_section = Path(filename).stem

    shared_metadata = {
        "source": filename,
        "session_id": session_id,
        "document_id": document_id,
        "upload_time": upload_time,
        "semantic_section": semantic_section,
    }
    for doc in documents:
        doc.metadata.update(shared_metadata)

    chunks = text_splitter.split_documents(documents)
    for index, chunk in enumerate(chunks):
        chunk.metadata.update(shared_metadata)
        chunk.metadata["chunk_id"] = f"{document_id}-chunk-{index:04d}"
    return chunks


def ingest_documents(
    *,
    file_bytes: bytes,
    filename: str,
    session_id: str = "default",
    content_type: str | None = None,
) -> dict:
    """Ingest one uploaded document: store it locally for the session, then
    index its chunks in Pinecone (dense) and BM25 (sparse)."""
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    session_key = sanitize_session_id(session_id)
    file_hash = calculate_content_hash(file_bytes)
    # Deterministic id from content so re-uploading the same file maps to a
    # stable id and overwrites existing vectors instead of duplicating them.
    document_id = f"doc_{file_hash[:16]}"
    upload_time = datetime.now(timezone.utc).isoformat()

    temp_path = _write_temp_file(filename, file_bytes)
    try:
        documents = load_document(temp_path)
    finally:
        if os.path.exists(temp_path):
            os.unlink(temp_path)

    if not documents:
        return {
            "chunks": 0,
            "pages": 0,
            "total_chunks": 0,
            "docs_indexed": 0,
            "session_id": session_key,
            "document_id": document_id,
            "indexed_files": [],
        }

    # Keep the raw file on local disk for the lifetime of the session; it is
    # removed when the session ends (explicit cleanup or the TTL sweep).
    stored = save_document(
        file_bytes=file_bytes,
        filename=filename,
        session_id=session_key,
        document_id=document_id,
    )

    chunks = _prepare_chunks(
        documents=documents,
        filename=filename,
        session_id=session_key,
        document_id=document_id,
        upload_time=upload_time,
    )
    vectors = get_embeddings().embed_documents([chunk.page_content for chunk in chunks])
    upsert_chunks(chunks, vectors)

    bm25_indexer = get_bm25_indexer(get_session_bm25_path(session_key))
    bm25_indexer.extend_and_rebuild(chunks)

    return {
        "chunks": len(chunks),
        "pages": len(documents),
        "total_chunks": len(chunks),
        "docs_indexed": 1,
        "session_id": session_key,
        "document_id": document_id,
        "indexed_files": [filename],
        "stored_uri": stored["uri"],
        "upload_time": upload_time,
    }


if __name__ == "__main__":
    raise SystemExit("Run ingestion via the FastAPI upload route.")
