import docx
import pytest

from app.services.ingestion import load_document


def test_load_txt(tmp_path):
    path = tmp_path / "note.txt"
    path.write_text("hybrid retrieval rocks", encoding="utf-8")

    docs = load_document(str(path))

    assert len(docs) == 1
    assert "hybrid retrieval" in docs[0].page_content


def test_load_docx_uses_python_docx(tmp_path):
    path = tmp_path / "memo.docx"
    document = docx.Document()
    document.add_paragraph("First paragraph about RAG.")
    document.add_paragraph("")  # blank paragraphs are skipped
    document.add_paragraph("Second paragraph about Qdrant.")
    document.save(str(path))

    docs = load_document(str(path))

    assert len(docs) == 1
    content = docs[0].page_content
    assert "First paragraph about RAG." in content
    assert "Second paragraph about Qdrant." in content


def test_unsupported_extension_raises(tmp_path):
    path = tmp_path / "image.png"
    path.write_bytes(b"\x89PNG\r\n")

    with pytest.raises(ValueError):
        load_document(str(path))
